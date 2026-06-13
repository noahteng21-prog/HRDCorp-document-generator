from docx import Document
from docx.shared import Pt

OUT_PATH = "templates/T3 PSMB_SBL_KHAS_T3_01 template.docx"

def make_template():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    doc.add_paragraph('FOR SBL-KHAS SCHEME ONLY', style='Title')
    doc.add_paragraph('')
    doc.add_paragraph('ATTENDANCE LIST', style='Heading 2')
    doc.add_paragraph('')
    doc.add_paragraph('Course Title : __________________')
    doc.add_paragraph('Dates of Training : __________________')
    doc.add_paragraph('')

    # Table with header
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    headers = ['No.', 'Name of Trainee(s)', 'Name of Employer(s)', 'NRIC', 'Citizenship', 'Sex', 'Signature*']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    # Insert loop start row
    start_row = table.add_row()
    start_row.cells[0].text = '{% for t in trainees %}'
    for i in range(1, 7):
        start_row.cells[i].text = ''

    # Template row
    tpl_row = table.add_row()
    tpl_row.cells[0].text = '{{ t.no }}'
    tpl_row.cells[1].text = '{{ t.name }}'
    tpl_row.cells[2].text = '{{ employer }}'
    tpl_row.cells[3].text = ''
    tpl_row.cells[4].text = ''
    tpl_row.cells[5].text = ''
    tpl_row.cells[6].text = ''

    # Insert loop end row
    end_row = table.add_row()
    end_row.cells[0].text = '{% endfor %}'
    for i in range(1, 7):
        end_row.cells[i].text = ''

    doc.add_paragraph('')
    doc.add_paragraph('I certify that all trainees listed above had fully attended the training.')

    doc.save(OUT_PATH)
    print(f"Template written to {OUT_PATH}")

if __name__ == '__main__':
    make_template()
